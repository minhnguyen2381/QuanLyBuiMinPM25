from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
REFERENCE_DOCX = Path(r"C:\Users\hasbr\Downloads\PPNCKH_LeHoan.docx")
RESULT_DIR = ROOT / "results" / "experiments" / "20260813_195553_pm25_hanoi_arima_rf_research"
OUTPUT_DOCX = ROOT / "docs" / "Bao_cao_tieu_luan_PM25_Phan_tich_du_lieu.docx"


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_page_layout(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.2)


def clear_headers_and_footers(doc: Document) -> None:
    """Xóa header/footer kế thừa từ file mẫu Word.

    Script dùng một file `.docx` có sẵn làm template định dạng. Nếu không xóa,
    header/footer cũ có thể xuất hiện trong báo cáo mới và gây nhầm lẫn.
    """
    for section in doc.sections:
        for part in [
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ]:
            for paragraph in part.paragraphs:
                paragraph.text = ""


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.first_line_indent = Cm(1.0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size, before, after, center in [
        ("Heading 1", 15, 12, 8, True),
        ("Heading 2", 13.5, 8, 4, False),
        ("Heading 3", 13, 6, 3, False),
    ]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.first_line_indent = Cm(0)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        st.paragraph_format.keep_with_next = True

    if "CaptionVN" not in styles:
        styles.add_style("CaptionVN", 1)
    caption = styles["CaptionVN"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(12)
    caption.font.italic = True
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str = "", style: str = "Normal", bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_centered(doc: Document, text: str, size: float = 13, bold: bool = False, spacing_after: int = 6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.first_line_indent = Cm(0)


def set_cell_text(cell, text: str, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11.5)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, rows: list[list], widths: list[float] | None = None) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell_text(table.cell(i, j), value, bold=(i == 0), center=(i == 0 or len(str(value)) < 12))
            if widths:
                table.cell(i, j).width = Cm(widths[j])
    doc.add_paragraph()


def add_compact_table(doc: Document, rows: list[list], widths: list[float] | None = None) -> None:
    """Thêm bảng phụ lục dạng gọn với chữ nhỏ hơn.

    Các bảng phụ lục thường nhiều dòng/cột hơn bảng chính. Hàm này giảm cỡ chữ
    và căn ô phù hợp để nội dung dài vẫn nằm trong trang Word.
    """
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or len(str(value)) < 10 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(10)
            run.bold = i == 0
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths:
                cell.width = Cm(widths[j])
    doc.add_paragraph()


def add_figure(doc: Document, image_path: Path, caption: str, width_cm: float = 15.5) -> None:
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = doc.add_paragraph(caption, style="CaptionVN")
    cap.paragraph_format.first_line_indent = Cm(0)


def df_rows(df: pd.DataFrame, columns: list[str], max_rows: int | None = None) -> list[list]:
    work = df.copy()
    if max_rows:
        work = work.head(max_rows)
    rows = [columns]
    for _, row in work.iterrows():
        rows.append([row.get(col, "") for col in columns])
    return rows


def format_float(value, digits=2):
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def add_cover(doc: Document) -> None:
    add_centered(doc, "BỘ CÔNG THƯƠNG\nĐẠI HỌC CÔNG NGHIỆP HÀ NỘI", 13, True, 22)
    add_centered(doc, "BÁO CÁO TIỂU LUẬN\nPHÂN TÍCH DỮ LIỆU", 16, True, 18)
    add_centered(
        doc,
        "NGHIÊN CỨU DỰ BÁO NỒNG ĐỘ BỤI MỊN PM2.5 TẠI HÀ NỘI\nBẰNG PHƯƠNG PHÁP ARIMA VÀ RANDOM FOREST",
        15,
        True,
        22,
    )

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Normal Table"
    info = [
        ("Giảng viên hướng dẫn :", "PGS. TS. Nguyễn Hữu Phấn"),
        ("Môn học :", "Phân tích dữ liệu"),
        ("Lớp :", "CH HTTT K15.3"),
        ("Mã lớp :", "20261ME7318001"),
        ("Học viên thực hiện :", "Lê Hoàn - 2025700249"),
    ]
    for i, (left, right) in enumerate(info):
        set_cell_text(table.cell(i, 0), left, bold=False)
        set_cell_text(table.cell(i, 1), right, bold=False)
    doc.add_paragraph("\n\n\n")
    add_centered(doc, "Hà Nội, Năm 2026", 13, False, 0)
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    figures = [
        ("Hình 2.1. Quy trình phân tích dữ liệu được thể hiện qua chuỗi PM2.5", "7"),
        ("Hình 3.1. Phân phối nồng độ PM2.5", "10"),
        ("Hình 3.2. Boxplot PM2.5 theo tháng", "10"),
        ("Hình 3.3. PM2.5 trung bình theo tháng qua các năm", "11"),
        ("Hình 3.4. Ma trận tương quan giữa các biến quan trắc", "12"),
        ("Hình 4.1. Đặc trưng quan trọng nhất của Random Forest", "14"),
        ("Hình 4.2. So sánh PM2.5 thực tế và dự báo của các mô hình", "15"),
        ("Hình 4.3. RMSE theo tháng của các mô hình", "15"),
        ("Hình 4.4. Phân phối sai số dự báo", "15"),
        ("Hình 4.5. Scatter giữa PM2.5 thực tế và PM2.5 dự báo", "16"),
    ]
    tables = [
        ("Bảng 2.1. Mô tả các nhóm biến sử dụng trong nghiên cứu", "7"),
        ("Bảng 3.1. Kết quả chất lượng dữ liệu", "9"),
        ("Bảng 3.2. Thống kê mô tả PM2.5 và một số biến ngoại sinh", "9"),
        ("Bảng 3.3. Thống kê PM2.5 theo mùa", "10"),
        ("Bảng 3.4. Tương quan giữa PM2.5 và các biến ngoại sinh", "11"),
        ("Bảng 4.1. Bảng so sánh kết quả mô hình", "13"),
        ("Bảng 4.2. Top đặc trưng quan trọng của Random Forest", "14"),
        ("Bảng 5.1. Tổng hợp kết quả và hàm ý", "17"),
    ]
    toc = [
        ("CHƯƠNG 1: MỞ ĐẦU", "5"),
        ("1.1. Lý do chọn đề tài", "5"),
        ("1.2. Mục tiêu và câu hỏi nghiên cứu", "5"),
        ("1.3. Đối tượng, phạm vi và ý nghĩa nghiên cứu", "6"),
        ("CHƯƠNG 2: CƠ SỞ LÝ LUẬN VÀ PHƯƠNG PHÁP", "6"),
        ("2.1. Tổng quan PM2.5 và bài toán dự báo", "6"),
        ("2.2. ARIMA và Random Forest trong phân tích dữ liệu", "7"),
        ("2.3. Quy trình phân tích dữ liệu", "7"),
        ("2.4. Các chỉ số đánh giá mô hình", "8"),
        ("2.5. Kiểm soát rò rỉ dữ liệu", "8"),
        ("CHƯƠNG 3: DỮ LIỆU VÀ PHÂN TÍCH KHÁM PHÁ", "9"),
        ("3.1. Nguồn dữ liệu và tiền xử lý", "9"),
        ("3.2. Thống kê mô tả và tính mùa vụ", "10"),
        ("3.3. Tương quan và đặc điểm biến ngoại sinh", "12"),
        ("CHƯƠNG 4: XÂY DỰNG, ĐÁNH GIÁ MÔ HÌNH", "13"),
        ("4.1. Thiết kế thực nghiệm", "13"),
        ("4.2. Kết quả ARIMA", "13"),
        ("4.3. Kết quả Random Forest", "14"),
        ("4.4. So sánh và phân tích sai số", "15"),
        ("CHƯƠNG 5: THẢO LUẬN VÀ KẾT LUẬN", "17"),
        ("5.1. Thảo luận kết quả", "17"),
        ("5.2. Hàm ý ứng dụng", "17"),
        ("5.3. Hạn chế và hướng phát triển", "18"),
        ("5.4. Kết luận", "18"),
        ("KẾT LUẬN", "19"),
        ("TÀI LIỆU THAM KHẢO", "20"),
        ("PHỤ LỤC", "21"),
    ]
    add_centered(doc, "DANH MỤC HÌNH ẢNH", 14, True, 8)
    add_table(doc, [["Tên hình", "Trang"], *figures], [14, 2])
    doc.add_page_break()
    add_centered(doc, "DANH MỤC BẢNG BIỂU", 14, True, 8)
    add_table(doc, [["Tên bảng", "Trang"], *tables], [14, 2])
    doc.add_page_break()
    add_centered(doc, "MỤC LỤC", 14, True, 8)
    add_table(doc, [["Nội dung", "Trang"], *toc], [14, 2])
    doc.add_page_break()


def add_method_note(doc: Document, title: str, paragraphs: list[str]) -> None:
    add_heading(doc, title, 3)
    for paragraph in paragraphs:
        add_paragraph(doc, paragraph)


def add_extended_discussion(doc: Document) -> None:
    add_heading(doc, "2.4. Các chỉ số đánh giá mô hình", 2)
    add_paragraph(doc, "Trong báo cáo này, RMSE được xem là chỉ số chính vì nó phạt mạnh các sai số lớn. Điều này phù hợp với bài toán PM2.5, bởi các ngày ô nhiễm tăng cao thường là những ngày có ý nghĩa cảnh báo lớn nhất. Nếu mô hình dự báo thấp đáng kể trong các ngày đỉnh ô nhiễm, hệ thống cảnh báo có thể đưa ra khuyến nghị muộn hoặc chưa đủ mạnh.")
    add_paragraph(doc, "MAE được sử dụng song song để đo sai số trung bình tuyệt đối theo đơn vị gốc µg/m³. MAPE và SMAPE giúp diễn giải sai số theo tỷ lệ phần trăm, nhưng cần thận trọng khi giá trị thực tế nhỏ. R2 cho biết tỷ lệ biến thiên được mô hình giải thích trên tập kiểm thử. Directional Accuracy đo khả năng dự báo đúng chiều tăng hoặc giảm giữa các ngày liên tiếp.")
    add_paragraph(doc, "Việc sử dụng nhiều chỉ số giúp tránh kết luận phiến diện. Một mô hình có RMSE thấp hơn nhưng Directional Accuracy không cao vẫn có thể dự báo mức độ tốt hơn nhưng chưa thật sự nắm bắt tốt chiều biến động ngắn hạn. Do đó, phần thảo luận cần nhìn đồng thời vào sai số độ lớn, sai số tỷ lệ và chiều biến động.")
    add_heading(doc, "2.5. Kiểm soát rò rỉ dữ liệu", 2)
    add_paragraph(doc, "Rò rỉ dữ liệu là rủi ro quan trọng trong dự báo chuỗi thời gian. Nếu đặc trưng đầu vào vô tình chứa giá trị của ngày cần dự báo hoặc các ngày tương lai, chỉ số đánh giá có thể đẹp nhưng không phản ánh năng lực dự báo thực tế. Vì vậy, toàn bộ đặc trưng lag, rolling và biến ngoại sinh trong Random Forest đều được dịch về quá khứ.")
    add_paragraph(doc, "Đối với ARIMA, quá trình chọn tham số chỉ sử dụng tập huấn luyện đến năm 2024. Tập kiểm thử năm 2025 chỉ được dùng để đánh giá sau cùng. Đối với Random Forest, việc tìm kiếm tham số sử dụng TimeSeriesSplit, tức là các fold vẫn tôn trọng thứ tự thời gian thay vì chia ngẫu nhiên.")
    add_paragraph(doc, "Nguyên tắc này đặc biệt quan trọng trong môn Phân tích dữ liệu vì kết quả mô hình không chỉ cần tốt về mặt số học, mà còn phải đúng về phương pháp. Một pipeline tránh rò rỉ dữ liệu giúp kết luận có giá trị hơn và có thể triển khai trong bối cảnh cảnh báo thực tế.")


def add_appendices(
    doc: Document,
    clean: pd.DataFrame,
    monthly: pd.DataFrame,
    yearly: pd.DataFrame,
    threshold: pd.DataFrame,
    rfimp: pd.DataFrame,
) -> None:
    doc.add_page_break()
    add_heading(doc, "PHỤ LỤC", 1)
    add_heading(doc, "Phụ lục 1. Mô tả bộ biến sau tiền xử lý", 2)
    add_paragraph(doc, "Phụ lục này trình bày vai trò của các biến trong bộ dữ liệu sạch đa biến. Các biến ngoại sinh không phải là biến mục tiêu, nhưng được sử dụng để bổ sung thông tin cho Random Forest sau khi đã dịch về quá khứ nhằm tránh rò rỉ dữ liệu.")
    data_dictionary = [
        ["Biến", "Nhóm", "Ý nghĩa sử dụng trong phân tích"],
        ["PM2.5", "Mục tiêu", "Nồng độ bụi mịn trung bình ngày cần dự báo"],
        ["AQI", "Chỉ số tổng hợp", "Phản ánh chất lượng không khí tổng hợp"],
        ["CO", "Chất ô nhiễm", "Biến liên quan đến đốt nhiên liệu và giao thông"],
        ["NO2", "Chất ô nhiễm", "Biến có liên hệ với phát thải giao thông và đô thị"],
        ["O3", "Chất ô nhiễm", "Biến ô nhiễm thứ cấp, liên quan đến phản ứng quang hóa"],
        ["PM10", "Chất ô nhiễm", "Bụi kích thước lớn hơn PM2.5, có tương quan đáng chú ý"],
        ["SO2", "Chất ô nhiễm", "Biến liên quan đến phát thải công nghiệp và đốt nhiên liệu"],
        ["Clouds", "Khí tượng", "Mức độ mây che phủ"],
        ["Precipitation", "Khí tượng", "Lượng mưa, có thể góp phần làm giảm bụi lơ lửng"],
        ["Pressure", "Khí tượng", "Áp suất khí quyển, liên quan điều kiện khuếch tán"],
        ["Relative Humidity", "Khí tượng", "Độ ẩm tương đối"],
        ["Temperature", "Khí tượng", "Nhiệt độ không khí"],
        ["UV Index", "Khí tượng", "Chỉ số tia cực tím"],
        ["Wind Speed", "Khí tượng", "Tốc độ gió, ảnh hưởng phát tán chất ô nhiễm"],
    ]
    add_compact_table(doc, data_dictionary, [3.5, 3.5, 9])

    add_heading(doc, "Phụ lục 2. Thống kê PM2.5 theo năm và theo tháng", 2)
    add_paragraph(doc, "Bảng thống kê theo năm giúp so sánh mức PM2.5 giữa các giai đoạn. Năm 2025 trong bộ dữ liệu hiện mới đến ngày 30/06/2025 nên không nên so sánh như một năm đầy đủ, nhưng vẫn có ý nghĩa khi đóng vai trò tập kiểm thử.")
    rows = [["Năm", "Số ngày", "Trung bình", "Trung vị", "Độ lệch chuẩn", "Nhỏ nhất", "Lớn nhất"]]
    for _, r in yearly.iterrows():
        rows.append([int(r["year"]), int(r["Số ngày"]), format_float(r["Trung bình"]), format_float(r["Trung vị"]), format_float(r["Độ lệch chuẩn"]), format_float(r["Nhỏ nhất"]), format_float(r["Lớn nhất"])])
    add_compact_table(doc, rows, [2, 2, 2.2, 2.2, 2.4, 2.2, 2.2])

    rows = [["Tháng", "Số ngày", "Trung bình", "Trung vị", "Độ lệch chuẩn", "Nhỏ nhất", "Lớn nhất"]]
    for _, r in monthly.iterrows():
        rows.append([int(r["month"]), int(r["Số ngày"]), format_float(r["Trung bình"]), format_float(r["Trung vị"]), format_float(r["Độ lệch chuẩn"]), format_float(r["Nhỏ nhất"]), format_float(r["Lớn nhất"])])
    add_compact_table(doc, rows, [2, 2, 2.2, 2.2, 2.4, 2.2, 2.2])

    add_heading(doc, "Phụ lục 3. Số ngày vượt ngưỡng tham khảo", 2)
    add_paragraph(doc, "Các ngưỡng trong bảng dưới đây không thay thế quy chuẩn pháp lý, mà được dùng như mốc tham khảo để mô tả mức độ xuất hiện của các ngày PM2.5 cao trong bộ dữ liệu.")
    rows = [["Ngưỡng PM2.5", "Số ngày vượt ngưỡng", "Tỷ lệ (%)"]]
    for _, r in threshold.iterrows():
        rows.append([format_float(r["Ngưỡng PM2.5"], 0), int(r["Số ngày vượt ngưỡng"]), format_float(r["Tỷ lệ (%)"], 2)])
    add_compact_table(doc, rows, [4, 5, 4])

    add_heading(doc, "Phụ lục 4. Danh sách đặc trưng Random Forest", 2)
    add_paragraph(doc, "Bảng này liệt kê toàn bộ đặc trưng được Random Forest đánh giá sau huấn luyện. Mức độ quan trọng là impurity-based importance của scikit-learn, nên dùng để tham khảo định hướng chứ không nên diễn giải như quan hệ nhân quả.")
    rows = [["STT", "Đặc trưng", "Mức độ quan trọng"]]
    for idx, (_, r) in enumerate(rfimp.iterrows(), start=1):
        rows.append([idx, r["Đặc trưng"], format_float(r["importance"], 5)])
    add_compact_table(doc, rows, [1.5, 9, 4])

    add_heading(doc, "Phụ lục 5. Chi tiết dự báo trên tập kiểm thử năm 2025", 2)
    add_paragraph(doc, "Bảng dưới đây là minh chứng chi tiết cho kết quả đánh giá mô hình. Các giá trị được lấy từ artifact của pipeline, căn theo ngày trong tập kiểm thử năm 2025.")
    actual = clean[clean.index.year >= 2025]["PM2.5"].copy()
    arima = pd.read_csv(RESULT_DIR / "predictions" / "arima_predictions.csv", index_col=0, parse_dates=True).iloc[:, 0]
    rf = pd.read_csv(RESULT_DIR / "predictions" / "random_forest_predictions.csv", index_col=0, parse_dates=True).iloc[:, 0]
    detail = pd.DataFrame({"Ngày": actual.index.strftime("%Y-%m-%d"), "Thực tế": actual.values})
    detail["ARIMA"] = arima.reindex(actual.index).values
    detail["Random Forest"] = rf.reindex(actual.index).values
    detail["Sai số tuyệt đối RF"] = (detail["Random Forest"] - detail["Thực tế"]).abs()
    rows = [["Ngày", "Thực tế", "ARIMA", "Random Forest", "|Sai số RF|"]]
    for _, r in detail.iterrows():
        rows.append([r["Ngày"], format_float(r["Thực tế"]), format_float(r["ARIMA"]), format_float(r["Random Forest"]), format_float(r["Sai số tuyệt đối RF"])])
    add_compact_table(doc, rows, [3, 2.5, 2.5, 3, 2.5])

    add_heading(doc, "Phụ lục 6. Cấu hình thực nghiệm và khả năng tái lập", 2)
    add_paragraph(doc, "Pipeline được điều khiển bởi các file YAML trong thư mục configs. Cách tổ chức này giúp thay đổi tham số mà không cần sửa trực tiếp mã nguồn, đồng thời giúp người đọc tái lập lại kết quả.")
    config_rows = [
        ["Thành phần", "Cấu hình chính"],
        ["Dữ liệu", "raw_dir=data/raw; clean_path=data/processed/clean_data_multivariate.csv; train_end_year=2024; test_start_year=2025"],
        ["ARIMA", "p_range=[0,1,2,3]; d_range=[0,1]; q_range=[0,1,2,3]; information_criterion=aic"],
        ["Random Forest", "lags=[1,2,3,7,14,21,30]; rolling_windows=[3,7,14,30]; include_exogenous=true"],
        ["Tuning", "RandomizedSearchCV; TimeSeriesSplit 5 folds; scoring=neg_root_mean_squared_error"],
        ["Đánh giá", "MAE, RMSE, MAPE, SMAPE, R2, DirectionalAccuracy"],
    ]
    add_compact_table(doc, config_rows, [4, 12])
    add_paragraph(doc, "Quy trình chạy lại gồm ba bước: cài đặt thư viện bằng `pip install -r requirements.txt`, kiểm tra dữ liệu gốc trong `data/raw`, sau đó chạy `python main.py`. Mỗi lần chạy tạo một thư mục kết quả riêng trong `results/experiments/<experiment_id>/`.")

    add_heading(doc, "Phụ lục 7. Checklist kiểm soát chất lượng phân tích", 2)
    checklist = [
        ["Nội dung kiểm soát", "Cách kiểm soát", "Tình trạng"],
        ["Dữ liệu gốc", "Kiểm tra đủ file CSV giai đoạn 2022-2025", "Đã thực hiện"],
        ["Tiền xử lý", "Chuẩn hóa ngày, ép kiểu số, tổng hợp theo ngày", "Đã thực hiện"],
        ["Dữ liệu thiếu", "Báo cáo missing sau tiền xử lý", "Đã thực hiện"],
        ["Rò rỉ dữ liệu", "Dùng shift cho lag, rolling và biến ngoại sinh", "Đã kiểm soát"],
        ["Chia train/test", "Chia theo thời gian, không chia ngẫu nhiên", "Đã thực hiện"],
        ["ARIMA", "Chọn p,d,q trên tập huấn luyện", "Đã thực hiện"],
        ["Random Forest", "Tuning bằng TimeSeriesSplit", "Đã thực hiện"],
        ["Đánh giá", "So sánh với mô hình tham chiếu", "Đã thực hiện"],
        ["Báo cáo", "Lưu bảng, hình và dự báo chi tiết", "Đã thực hiện"],
    ]
    add_compact_table(doc, checklist, [4.5, 7, 4])

    add_heading(doc, "Phụ lục 8. Nhật ký quan sát PM2.5 theo ngày năm 2024", 2)
    add_paragraph(doc, "Bảng nhật ký quan sát trình bày dữ liệu theo ngày của một năm đầy đủ trong tập huấn luyện. Phụ lục này giúp người đọc kiểm tra trực tiếp nền dữ liệu đầu vào trước khi mô hình được huấn luyện, đồng thời thể hiện rõ tính liên tục của chuỗi thời gian.")
    observed_2024 = clean[clean.index.year == 2024].copy()
    observed_rows = [["Ngày", "PM2.5", "AQI", "PM10", "Nhiệt độ", "Độ ẩm", "Gió"]]
    for idx, r in observed_2024.iterrows():
        observed_rows.append([
            idx.strftime("%Y-%m-%d"),
            format_float(r.get("PM2.5")),
            format_float(r.get("AQI")),
            format_float(r.get("PM10")),
            format_float(r.get("Temperature")),
            format_float(r.get("Relative Humidity")),
            format_float(r.get("Wind Speed")),
        ])
    add_compact_table(doc, observed_rows, [2.8, 2, 2, 2, 2.2, 2.2, 2])

    add_heading(doc, "Phụ lục 9. Phân tầng sai số Random Forest theo mức PM2.5 thực tế", 2)
    add_paragraph(doc, "Bảng phân tầng sai số giúp đánh giá mô hình theo các mức ô nhiễm khác nhau. Nếu sai số tăng mạnh ở nhóm PM2.5 cao, hệ thống cảnh báo cần bổ sung cơ chế giám sát riêng cho các ngày ô nhiễm nặng.")
    bins = [0, 25, 50, 75, 100, float("inf")]
    labels = ["<=25", "25-50", "50-75", "75-100", ">100"]
    detail["Nhóm PM2.5"] = pd.cut(detail["Thực tế"], bins=bins, labels=labels, include_lowest=True)
    grouped_error = detail.groupby("Nhóm PM2.5", observed=False).agg(
        So_ngay=("Thực tế", "count"),
        PM25_tb=("Thực tế", "mean"),
        Sai_so_RF_tb=("Sai số tuyệt đối RF", "mean"),
        Sai_so_RF_trung_vi=("Sai số tuyệt đối RF", "median"),
        Sai_so_RF_lon_nhat=("Sai số tuyệt đối RF", "max"),
    ).reset_index()
    error_rows = [["Nhóm PM2.5", "Số ngày", "PM2.5 TB", "|Sai số RF| TB", "Trung vị", "Lớn nhất"]]
    for _, r in grouped_error.iterrows():
        error_rows.append([
            r["Nhóm PM2.5"],
            int(r["So_ngay"]),
            format_float(r["PM25_tb"]),
            format_float(r["Sai_so_RF_tb"]),
            format_float(r["Sai_so_RF_trung_vi"]),
            format_float(r["Sai_so_RF_lon_nhat"]),
        ])
    add_compact_table(doc, error_rows, [3, 2, 2.6, 3, 2.4, 2.4])

    add_paragraph(doc, "Nhìn vào phụ lục sai số, có thể thấy việc đánh giá mô hình không chỉ dựa vào một chỉ số trung bình toàn tập. Trong ứng dụng cảnh báo PM2.5, các ngày có nồng độ cao thường quan trọng hơn về mặt quản lý rủi ro, vì vậy cần tiếp tục theo dõi sai số theo từng nhóm nồng độ trong các nghiên cứu sau.")


def build_report() -> None:
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(REFERENCE_DOCX, OUTPUT_DOCX)
    doc = Document(OUTPUT_DOCX)
    clear_body(doc)
    set_page_layout(doc)
    configure_styles(doc)
    clear_headers_and_footers(doc)

    comparison = pd.read_csv(RESULT_DIR / "model_comparison.csv")
    clean = pd.read_csv(ROOT / "data" / "processed" / "clean_data_multivariate.csv", index_col=0, parse_dates=True)
    desc = pd.read_csv(RESULT_DIR / "tables" / "thong_ke_mo_ta.csv").rename(columns={"Unnamed: 0": "Biến"})
    monthly = pd.read_csv(RESULT_DIR / "tables" / "thong_ke_pm25_theo_monthly.csv")
    seasonal = pd.read_csv(RESULT_DIR / "tables" / "thong_ke_pm25_theo_seasonal.csv")
    yearly = pd.read_csv(RESULT_DIR / "tables" / "thong_ke_pm25_theo_yearly.csv")
    corr = pd.read_csv(RESULT_DIR / "tables" / "tuong_quan_pm25_bien_ngoai_sinh.csv").rename(columns={"Unnamed: 0": "Biến"})
    threshold = pd.read_csv(RESULT_DIR / "tables" / "so_ngay_vuot_nguong_tham_khao.csv")
    rfimp = pd.read_csv(RESULT_DIR / "random_forest_feature_importance.csv").rename(columns={"Unnamed: 0": "Đặc trưng"})
    quality = pd.read_csv(RESULT_DIR / "data_quality_report.csv")

    fig_dir = RESULT_DIR / "figures"

    add_cover(doc)
    add_front_matter(doc)

    add_heading(doc, "CHƯƠNG 1: MỞ ĐẦU", 1)
    add_heading(doc, "1.1. Lý do chọn đề tài", 2)
    add_paragraph(doc, "Ô nhiễm không khí là một trong những vấn đề môi trường đô thị đáng quan tâm tại các thành phố lớn. Trong đó, bụi mịn PM2.5 có kích thước nhỏ, có khả năng tồn tại lâu trong không khí và ảnh hưởng trực tiếp đến sức khỏe con người. Tại Hà Nội, nồng độ PM2.5 thường biến động mạnh theo thời gian, chịu tác động đồng thời của điều kiện khí tượng, giao thông, xây dựng, hoạt động dân sinh và các nguồn phát thải khác.")
    add_paragraph(doc, "Đối với môn học Phân tích dữ liệu, đề tài dự báo PM2.5 có tính phù hợp vì dữ liệu quan trắc có cấu trúc chuỗi thời gian, có nhiều biến ngoại sinh và có khả năng triển khai đầy đủ quy trình phân tích: thu thập dữ liệu, tiền xử lý, mô tả, trực quan hóa, xây dựng mô hình, đánh giá và diễn giải kết quả. Đề tài không chỉ yêu cầu chạy mô hình, mà còn yêu cầu giải thích dữ liệu và đánh giá độ tin cậy của kết quả dự báo.")
    add_paragraph(doc, "Từ góc độ Hệ thống thông tin, dự báo PM2.5 có thể được xem là một thành phần trong hệ thống hỗ trợ ra quyết định về môi trường đô thị. Nếu được triển khai trong thực tế, kết quả dự báo có thể phục vụ cảnh báo sớm, lập kế hoạch truyền thông sức khỏe, hỗ trợ người dân điều chỉnh hoạt động ngoài trời và cung cấp căn cứ dữ liệu cho cơ quan quản lý.")
    add_heading(doc, "1.2. Mục tiêu và câu hỏi nghiên cứu", 2)
    add_paragraph(doc, "Mục tiêu tổng quát của tiểu luận là phân tích đặc điểm biến động PM2.5 tại Hà Nội giai đoạn 2022-2025 và xây dựng mô hình dự báo ngắn hạn bằng ARIMA và Random Forest. Từ đó, bài báo cáo so sánh năng lực dự báo của hai hướng tiếp cận: mô hình thống kê chuỗi thời gian và mô hình học máy dựa trên đặc trưng.")
    add_paragraph(doc, "Các mục tiêu cụ thể gồm: (1) mô tả và làm sạch bộ dữ liệu quan trắc không khí; (2) phân tích xu hướng, mùa vụ, phân phối và tương quan giữa PM2.5 với các biến ngoại sinh; (3) xây dựng mô hình ARIMA và Random Forest; (4) đánh giá mô hình bằng MAE, RMSE, MAPE, SMAPE, R2 và Directional Accuracy; (5) rút ra hàm ý ứng dụng cho bài toán cảnh báo chất lượng không khí.")
    add_paragraph(doc, "Các câu hỏi nghiên cứu chính gồm: PM2.5 tại Hà Nội biến động như thế nào theo tháng, mùa và năm? Những biến ngoại sinh nào có tương quan đáng chú ý với PM2.5? ARIMA và Random Forest đạt kết quả dự báo ra sao trên tập kiểm thử năm 2025? Mô hình nào phù hợp hơn trong phạm vi dữ liệu và thiết kế thực nghiệm của tiểu luận?")
    add_heading(doc, "1.3. Đối tượng, phạm vi và ý nghĩa nghiên cứu", 2)
    add_paragraph(doc, "Đối tượng phân tích là nồng độ bụi mịn PM2.5 trung bình ngày tại Hà Nội. Dữ liệu sử dụng trong báo cáo được tổng hợp từ bộ dữ liệu chất lượng không khí giai đoạn 2022-2025, bao gồm PM2.5 và các biến liên quan như AQI, CO, NO2, O3, PM10, SO2, mây, mưa, áp suất, độ ẩm, nhiệt độ, UV Index và tốc độ gió.")
    add_paragraph(doc, "Phạm vi thời gian sau tiền xử lý kéo dài từ ngày 13/01/2022 đến ngày 30/06/2025. Tập huấn luyện được xác định đến hết năm 2024, trong khi tập kiểm thử sử dụng dữ liệu từ năm 2025. Cách chia này phù hợp với dữ liệu chuỗi thời gian vì tránh việc sử dụng thông tin tương lai cho huấn luyện.")
    add_paragraph(doc, "Ý nghĩa khoa học của đề tài nằm ở việc vận dụng quy trình phân tích dữ liệu vào một bài toán thực tiễn, có cả yếu tố mô tả, dự báo và diễn giải. Ý nghĩa thực tiễn nằm ở khả năng đề xuất một quy trình dữ liệu có thể tích hợp vào hệ thống cảnh báo chất lượng không khí trong tương lai.")

    add_heading(doc, "CHƯƠNG 2: CƠ SỞ LÝ LUẬN VÀ PHƯƠNG PHÁP", 1)
    add_heading(doc, "2.1. Tổng quan PM2.5 và bài toán dự báo", 2)
    add_paragraph(doc, "PM2.5 là bụi mịn có đường kính khí động học nhỏ hơn hoặc bằng 2,5 micromet. Do kích thước nhỏ, PM2.5 có thể xâm nhập sâu vào hệ hô hấp và liên quan đến nhiều nguy cơ sức khỏe. Trong phân tích dữ liệu môi trường, PM2.5 thường được xem là một biến mục tiêu quan trọng vì phản ánh mức độ ô nhiễm không khí ở quy mô đô thị.")
    add_paragraph(doc, "Bài toán dự báo PM2.5 là bài toán chuỗi thời gian có tính bất định cao. Giá trị PM2.5 ngày hôm nay thường phụ thuộc vào giá trị các ngày trước, nhưng cũng chịu ảnh hưởng của điều kiện khí tượng và các chất ô nhiễm khác. Vì vậy, chỉ dùng một mô hình đơn biến có thể chưa khai thác hết thông tin trong dữ liệu.")
    add_heading(doc, "2.2. ARIMA và Random Forest trong phân tích dữ liệu", 2)
    add_paragraph(doc, "ARIMA là mô hình thống kê cổ điển cho chuỗi thời gian, trong đó p thể hiện bậc tự hồi quy, d thể hiện số lần sai phân và q thể hiện bậc trung bình trượt. Ưu điểm của ARIMA là có cơ sở thống kê rõ ràng, dễ giải thích và phù hợp với dự báo ngắn hạn. Hạn chế của ARIMA là khó biểu diễn quan hệ phi tuyến và thường cần xử lý kỹ tính dừng của chuỗi.")
    add_paragraph(doc, "Random Forest Regression là mô hình học máy dựa trên tập hợp nhiều cây quyết định. Đối với dữ liệu chuỗi thời gian, Random Forest không tự hiểu thứ tự thời gian, nên cần chuyển chuỗi thành bài toán học có giám sát bằng các đặc trưng như lag, rolling mean, rolling standard deviation, đặc trưng lịch và biến ngoại sinh đã dịch về quá khứ. Ưu điểm của Random Forest là mô hình hóa được quan hệ phi tuyến và cung cấp thông tin về mức độ quan trọng của đặc trưng.")
    add_heading(doc, "2.3. Quy trình phân tích dữ liệu", 2)
    add_paragraph(doc, "Quy trình nghiên cứu được tổ chức theo sáu bước: thu thập dữ liệu, tiền xử lý, phân tích khám phá, tạo đặc trưng, huấn luyện mô hình và đánh giá kết quả. Toàn bộ quy trình được cài đặt thành pipeline Python có cấu hình YAML, giúp tái lập kết quả và lưu artifact theo từng lần chạy.")
    add_figure(doc, fig_dir / "01_dien_bien_pm25.png", "Hình 2.1. Quy trình phân tích dữ liệu được thể hiện thông qua chuỗi kết quả PM2.5 toàn giai đoạn", 14.5)
    rows = [
        ["Nhóm biến", "Biến tiêu biểu", "Vai trò trong nghiên cứu"],
        ["Biến mục tiêu", "PM2.5", "Giá trị cần dự báo"],
        ["Chất ô nhiễm", "AQI, PM10, NO2, CO, O3, SO2", "Biến ngoại sinh hỗ trợ Random Forest"],
        ["Khí tượng", "Temperature, Relative Humidity, Wind Speed, Pressure, Clouds, Precipitation", "Giải thích bối cảnh biến động PM2.5"],
        ["Thời gian", "Tháng, quý, mùa, ngày trong tuần", "Biểu diễn mùa vụ và chu kỳ"],
        ["Đặc trưng trễ", "pm25_lag_1, pm25_lag_7, rolling mean", "Khai thác phụ thuộc quá khứ"],
    ]
    add_paragraph(doc, "Bảng 2.1. Mô tả các nhóm biến sử dụng trong nghiên cứu", style="CaptionVN")
    add_table(doc, rows, [4, 6, 6])
    add_extended_discussion(doc)

    doc.add_page_break()
    add_heading(doc, "CHƯƠNG 3: DỮ LIỆU VÀ PHÂN TÍCH KHÁM PHÁ", 1)
    add_heading(doc, "3.1. Nguồn dữ liệu và tiền xử lý", 2)
    add_paragraph(doc, "Dữ liệu thô gồm 30.341 bản ghi quan trắc theo thời gian. Sau khi chuẩn hóa cột thời gian, ép kiểu số, loại bỏ bản ghi không hợp lệ, tổng hợp theo ngày và nội suy các khoảng thiếu ngắn, bộ dữ liệu sạch còn 1.265 quan sát ngày. Dữ liệu sau xử lý không còn thiếu giá trị PM2.5 và giữ lại 14 biến số phục vụ phân tích đa biến.")
    quality_rows = [
        ["Chỉ tiêu", "Giá trị"],
        ["Số bản ghi thô", "30.341"],
        ["Số ngày sau tiền xử lý", "1.265"],
        ["Thời gian bắt đầu", "13/01/2022"],
        ["Thời gian kết thúc", "30/06/2025"],
        ["Giá trị PM2.5 thiếu sau xử lý", "0"],
        ["Số biến sau xử lý", "14"],
    ]
    add_paragraph(doc, "Bảng 3.1. Kết quả chất lượng dữ liệu", style="CaptionVN")
    add_table(doc, quality_rows, [7, 7])
    add_heading(doc, "3.2. Thống kê mô tả và tính mùa vụ", 2)
    add_paragraph(doc, "Nồng độ PM2.5 trung bình toàn kỳ đạt 50,77 µg/m³, trung vị 41,50 µg/m³, độ lệch chuẩn 32,69 và giá trị lớn nhất 261,01 µg/m³. Phân phối PM2.5 lệch phải rõ rệt với skewness 2,09, cho thấy tồn tại một số giai đoạn ô nhiễm tăng cao bất thường.")
    desc_sel = desc[desc["Biến"].isin(["PM2.5", "AQI", "PM10", "NO2", "Temperature", "Relative Humidity", "Wind Speed", "Pressure"])]
    rows = [["Biến", "Trung bình", "Trung vị", "Độ lệch chuẩn", "Nhỏ nhất", "Lớn nhất"]]
    for _, r in desc_sel.iterrows():
        rows.append([r["Biến"], format_float(r["mean"]), format_float(r["median"]), format_float(r["std"]), format_float(r["min"]), format_float(r["max"])])
    add_paragraph(doc, "Bảng 3.2. Thống kê mô tả PM2.5 và một số biến ngoại sinh", style="CaptionVN")
    add_table(doc, rows, [3, 2.5, 2.5, 2.8, 2.3, 2.5])
    add_figure(doc, fig_dir / "02_phan_phoi_pm25.png", "Hình 3.1. Phân phối nồng độ PM2.5 sau tiền xử lý", 14)
    add_paragraph(doc, "Kết quả theo mùa cho thấy PM2.5 cao nhất vào mùa đông với trung bình 67,30 µg/m³, tiếp đến là mùa thu 54,84 µg/m³ và mùa xuân 50,98 µg/m³. Mùa hè có mức thấp nhất với trung bình 29,71 µg/m³. Sự khác biệt này phù hợp với đặc điểm khí tượng đô thị, khi mùa đông thường có điều kiện khuếch tán kém hơn.")
    rows = [["Mùa", "Số ngày", "Trung bình", "Trung vị", "Độ lệch chuẩn", "Lớn nhất"]]
    for _, r in seasonal.iterrows():
        rows.append([r["season"], int(r["Số ngày"]), format_float(r["Trung bình"]), format_float(r["Trung vị"]), format_float(r["Độ lệch chuẩn"]), format_float(r["Lớn nhất"])])
    add_paragraph(doc, "Bảng 3.3. Thống kê PM2.5 theo mùa", style="CaptionVN")
    add_table(doc, rows, [4.2, 2, 2.5, 2.5, 2.8, 2.4])
    add_figure(doc, fig_dir / "03_boxplot_pm25_theo_thang.png", "Hình 3.2. Boxplot PM2.5 theo tháng", 14)
    add_figure(doc, fig_dir / "04_trung_binh_pm25_theo_thang_nam.png", "Hình 3.3. PM2.5 trung bình theo tháng qua các năm", 15)
    add_heading(doc, "3.3. Tương quan và đặc điểm biến ngoại sinh", 2)
    add_paragraph(doc, "Tương quan giữa PM2.5 với AQI đạt 0,8590 và với PM10 đạt 0,6648. Đây là hai quan hệ mạnh nhất trong bộ dữ liệu. Áp suất có tương quan dương 0,4290, NO2 có tương quan dương 0,4206, trong khi nhiệt độ có tương quan âm -0,3736. Điều này cho thấy mô hình học máy có cơ sở để khai thác thêm các biến ngoại sinh thay vì chỉ dùng lịch sử PM2.5.")
    rows = [["Biến", "Tương quan với PM2.5"]]
    for _, r in corr.head(8).iterrows():
        rows.append([r["Biến"], format_float(r["Tương quan với PM2.5"], 4)])
    add_paragraph(doc, "Bảng 3.4. Tương quan giữa PM2.5 và các biến ngoại sinh", style="CaptionVN")
    add_table(doc, rows, [7, 5])
    add_figure(doc, fig_dir / "05_heatmap_tuong_quan.png", "Hình 3.4. Ma trận tương quan giữa các biến quan trắc", 14.5)
    add_paragraph(doc, "Bảng số ngày vượt ngưỡng cho thấy 98,02% số ngày có PM2.5 lớn hơn 15 µg/m³, 83,24% số ngày lớn hơn 25 µg/m³ và 37,15% số ngày lớn hơn 50 µg/m³. Điều này củng cố tính cấp thiết của việc theo dõi và dự báo PM2.5 trong môi trường đô thị.")

    doc.add_page_break()
    add_heading(doc, "CHƯƠNG 4: XÂY DỰNG, ĐÁNH GIÁ MÔ HÌNH", 1)
    add_heading(doc, "4.1. Thiết kế thực nghiệm", 2)
    add_paragraph(doc, "Thực nghiệm được thiết kế theo nguyên tắc dữ liệu chuỗi thời gian. Tập huấn luyện gồm dữ liệu từ năm 2022 đến năm 2024; tập kiểm thử gồm dữ liệu từ năm 2025. Mô hình tham chiếu (Naive) sử dụng giá trị ngày trước làm dự báo cho ngày hiện tại. Tham chiếu mùa vụ 7 ngày sử dụng giá trị cùng thứ trong tuần trước. Hai mốc này giúp đánh giá liệu ARIMA và Random Forest có tạo ra giá trị dự báo thực sự hay không.")
    add_paragraph(doc, "ARIMA được chọn tham số bằng grid-search trên p, d, q theo tiêu chí AIC. Random Forest được huấn luyện trên các đặc trưng lag, rolling, lịch thời gian, đặc trưng chu kỳ và biến ngoại sinh đã shift về quá khứ. Việc shift biến ngoại sinh giúp tránh rò rỉ dữ liệu tương lai.")
    add_heading(doc, "4.2. Kết quả mô hình ARIMA", 2)
    add_paragraph(doc, "Kết quả chọn tham số cho thấy mô hình ARIMA(3,1,1) là cấu hình được lựa chọn theo tiêu chí AIC. Trên tập kiểm thử năm 2025, ARIMA đạt MAE 15,7073, RMSE 21,2486, MAPE 30,3311%, SMAPE 26,0787%, R2 đạt 0,4949 và cải thiện RMSE 5,12% so với mô hình tham chiếu. Kết quả này cho thấy ARIMA khai thác được một phần cấu trúc chuỗi thời gian, nhưng vẫn còn hạn chế trong việc mô hình hóa biến động phi tuyến.")
    add_heading(doc, "4.3. Kết quả mô hình Random Forest", 2)
    add_paragraph(doc, "Random Forest đạt kết quả tốt nhất trong các mô hình được so sánh. Mô hình đạt MAE 14,3909, RMSE 19,2163, MAPE 26,3409%, SMAPE 24,1010%, R2 đạt 0,5869 và cải thiện RMSE 14,19% so với mô hình tham chiếu. Điều này cho thấy việc bổ sung đặc trưng ngoại sinh và đặc trưng trễ giúp mô hình học máy nắm bắt tốt hơn sự biến động của PM2.5.")
    rows = [["Mô hình", "MAE", "RMSE", "MAPE", "SMAPE", "R2", "Cải thiện RMSE (%)"]]
    for _, r in comparison.iterrows():
        rows.append([r["Model"], format_float(r["MAE"], 4), format_float(r["RMSE"], 4), format_float(r["MAPE"], 4), format_float(r["SMAPE"], 4), format_float(r["R2"], 4), format_float(r["Cải thiện RMSE so với mô hình tham chiếu (%)"], 2)])
    add_paragraph(doc, "Bảng 4.1. Bảng so sánh kết quả mô hình", style="CaptionVN")
    add_table(doc, rows, [4.3, 2, 2, 2, 2, 1.8, 2.8])
    rows = [["Đặc trưng", "Mức độ quan trọng"]]
    for _, r in rfimp.head(10).iterrows():
        rows.append([r["Đặc trưng"], format_float(r["importance"], 4)])
    add_paragraph(doc, "Bảng 4.2. Top đặc trưng quan trọng của Random Forest", style="CaptionVN")
    add_table(doc, rows, [9, 4])
    add_figure(doc, fig_dir / "06_dac_trung_quan_trong_random_forest.png", "Hình 4.1. Đặc trưng quan trọng nhất của Random Forest", 14)
    add_heading(doc, "4.4. So sánh và phân tích sai số", 2)
    add_paragraph(doc, "So với ARIMA, Random Forest có RMSE thấp hơn khoảng 2,03 µg/m³ và R2 cao hơn. Kết quả này cho thấy trong phạm vi dữ liệu hiện tại, mô hình học máy đa biến có lợi thế hơn mô hình chuỗi thời gian đơn biến. Tuy nhiên, Directional Accuracy của Random Forest chỉ đạt 0,5167, nghĩa là khả năng dự đoán đúng chiều tăng/giảm vẫn còn khiêm tốn.")
    add_figure(doc, RESULT_DIR / "forecast_comparison.png", "Hình 4.2. So sánh PM2.5 thực tế và dự báo của các mô hình", 15)
    add_figure(doc, fig_dir / "07_rmse_theo_thang.png", "Hình 4.3. RMSE theo tháng của các mô hình", 15)
    add_figure(doc, fig_dir / "08_phan_phoi_sai_so_du_bao.png", "Hình 4.4. Phân phối sai số dự báo", 15)
    add_figure(doc, fig_dir / "09_thuc_te_va_du_bao_scatter.png", "Hình 4.5. Scatter giữa PM2.5 thực tế và PM2.5 dự báo", 13)

    doc.add_page_break()
    add_heading(doc, "CHƯƠNG 5: THẢO LUẬN VÀ KẾT LUẬN", 1)
    add_heading(doc, "5.1. Thảo luận kết quả", 2)
    add_paragraph(doc, "Kết quả phân tích khám phá cho thấy PM2.5 tại Hà Nội có tính mùa vụ rõ rệt. Các tháng cuối năm và đầu năm thường có mức PM2.5 cao hơn, trong khi mùa hè có mức thấp hơn. Phân phối lệch phải và các giá trị cực đại cho thấy dữ liệu có những đợt ô nhiễm cao, gây khó khăn cho dự báo bằng các mô hình tuyến tính đơn giản.")
    add_paragraph(doc, "Tương quan mạnh giữa PM2.5 với AQI và PM10 là hợp lý vì các biến này cùng phản ánh mức độ ô nhiễm không khí. Tương quan với áp suất, NO2 và nhiệt độ cũng gợi ý rằng bối cảnh khí tượng và các chất ô nhiễm khác có vai trò đáng kể. Do đó, việc Random Forest sử dụng biến ngoại sinh là một hướng tiếp cận phù hợp trong môn Phân tích dữ liệu.")
    add_heading(doc, "5.2. Hàm ý ứng dụng", 2)
    add_paragraph(doc, "Nếu triển khai thành hệ thống thông tin hỗ trợ cảnh báo, pipeline có thể được tổ chức thành các thành phần: thu thập dữ liệu quan trắc, kiểm tra chất lượng, cập nhật dữ liệu sạch, chạy mô hình dự báo, lưu kết quả, trực quan hóa trên dashboard và phát cảnh báo khi PM2.5 vượt ngưỡng. Mô hình Random Forest có thể được dùng làm mô hình dự báo chính trong giai đoạn đầu, trong khi ARIMA đóng vai trò mô hình thống kê đối chứng.")
    rows = [
        ["Nội dung", "Kết quả chính", "Hàm ý"],
        ["Dữ liệu", "1.265 ngày, không thiếu PM2.5 sau xử lý", "Đủ điều kiện cho phân tích chuỗi ngày"],
        ["EDA", "PM2.5 trung bình 50,77 µg/m³; mùa đông cao nhất", "Cần chú ý yếu tố mùa vụ trong cảnh báo"],
        ["Tương quan", "AQI, PM10, áp suất, NO2 liên quan đáng chú ý", "Nên khai thác biến ngoại sinh"],
        ["ARIMA", "Cải thiện RMSE 5,12%", "Phù hợp làm mô hình thống kê đối chứng"],
        ["Random Forest", "Cải thiện RMSE 14,19%", "Phù hợp hơn cho dự báo đa biến ngắn hạn"],
    ]
    add_paragraph(doc, "Bảng 5.1. Tổng hợp kết quả và hàm ý", style="CaptionVN")
    add_table(doc, rows, [4, 6, 6])
    add_heading(doc, "5.3. Hạn chế và hướng phát triển", 2)
    add_paragraph(doc, "Hạn chế thứ nhất là dữ liệu chỉ đại diện cho nguồn dữ liệu quan trắc đã có, chưa phân tích chi tiết theo từng trạm hoặc từng khu vực nhỏ trong thành phố. Hạn chế thứ hai là nghiên cứu mới dự báo ở mức ngày, chưa mở rộng sang dự báo theo giờ. Hạn chế thứ ba là mô hình chưa kiểm định đầy đủ sự khác biệt sai số bằng các kiểm định thống kê như Diebold-Mariano.")
    add_paragraph(doc, "Hướng phát triển tiếp theo gồm: bổ sung dữ liệu nhiều trạm, mở rộng dự báo theo giờ, thử SARIMAX để đưa biến ngoại sinh vào họ mô hình ARIMA, bổ sung permutation importance để giải thích Random Forest, và xây dựng dashboard trực quan hóa kết quả dự báo theo thời gian.")
    add_heading(doc, "5.4. Kết luận", 2)
    add_paragraph(doc, "Tiểu luận đã thực hiện đầy đủ quy trình phân tích dữ liệu cho bài toán dự báo PM2.5 tại Hà Nội. Kết quả cho thấy PM2.5 có biến động mùa vụ rõ ràng, phân phối lệch phải và có quan hệ đáng chú ý với nhiều biến ngoại sinh. Trong hai mô hình chính, Random Forest cho kết quả tốt hơn ARIMA trên tập kiểm thử năm 2025, với RMSE 19,2163 và cải thiện 14,19% so với mô hình tham chiếu.")
    add_paragraph(doc, "Trong phạm vi môn Phân tích dữ liệu, kết quả quan trọng không chỉ là mô hình nào có chỉ số tốt hơn, mà là toàn bộ quy trình từ dữ liệu đến kết luận được tổ chức có hệ thống, có khả năng tái lập và có diễn giải gắn với bối cảnh thực tế. Đây là nền tảng để phát triển các hệ thống hỗ trợ ra quyết định về chất lượng không khí trong tương lai.")

    doc.add_page_break()
    add_heading(doc, "KẾT LUẬN", 1)
    add_paragraph(doc, "Báo cáo đã hoàn thành mục tiêu nghiên cứu dự báo nồng độ bụi mịn PM2.5 tại Hà Nội bằng ARIMA và Random Forest trên bộ dữ liệu quan trắc giai đoạn 2022-2025. Toàn bộ quy trình được triển khai theo logic của môn Phân tích dữ liệu, từ tiền xử lý, khám phá dữ liệu, xây dựng đặc trưng, huấn luyện mô hình đến đánh giá và trình bày kết quả.")
    add_paragraph(doc, "Về dữ liệu, bộ dữ liệu sau tiền xử lý gồm 1.265 quan sát ngày, không còn thiếu giá trị PM2.5 và giữ lại nhiều biến ngoại sinh có ý nghĩa. Phân tích khám phá cho thấy PM2.5 tại Hà Nội có tính mùa vụ rõ, mùa đông cao nhất và mùa hè thấp nhất. Phân phối PM2.5 lệch phải, thể hiện sự tồn tại của các đợt ô nhiễm tăng cao.")
    add_paragraph(doc, "Về mô hình, ARIMA(3,1,1) cải thiện RMSE 5,12% so với mô hình tham chiếu, trong khi Random Forest cải thiện 14,19%. Kết quả này cho thấy Random Forest phù hợp hơn trong bối cảnh có nhiều biến ngoại sinh và quan hệ phi tuyến. Tuy nhiên, kết luận này chỉ đúng trong phạm vi bộ dữ liệu và thiết kế thực nghiệm của báo cáo, không nên khái quát cho mọi thành phố hoặc mọi giai đoạn.")
    add_paragraph(doc, "Về ý nghĩa học phần, tiểu luận nhấn mạnh rằng phân tích dữ liệu không dừng ở việc chạy thuật toán. Giá trị của bài phân tích nằm ở khả năng đặt vấn đề, hiểu dữ liệu, kiểm soát rò rỉ dữ liệu, chọn mô hình phù hợp, đánh giá bằng nhiều tiêu chí và diễn giải kết quả theo bối cảnh thực tiễn.")

    doc.add_page_break()
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    refs = [
        "Box, G. E. P., Jenkins, G. M., Reinsel, G. C., & Ljung, G. M. (2015). Time Series Analysis: Forecasting and Control.",
        "Breiman, L. (2001). Random Forests. Machine Learning, 45, 5-32.",
        "Hyndman, R. J., & Athanasopoulos, G. (2021). Forecasting: Principles and Practice.",
        "World Health Organization. (2021). WHO global air quality guidelines.",
        "Scikit-learn documentation. RandomForestRegressor.",
        "Statsmodels documentation. ARIMA models.",
        "Bộ dữ liệu Kaggle: phungdinhdat/aqi-in-hanoi-2022-2025.",
    ]
    for ref in refs:
        p = add_paragraph(doc, ref)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)

    add_appendices(doc, clean, monthly, yearly, threshold, rfimp)

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_report()
